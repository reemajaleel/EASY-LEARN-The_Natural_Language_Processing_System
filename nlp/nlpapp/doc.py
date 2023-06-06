# import pydoc
# document = pydoc('C:\\Users\\Reema\\PycharmProjects\\nlp\\media\\20230311134840Word.docx')
# print(document.paragraphs)
# from datetime import  datetime
#
# # date=datetime.now().strftime("%Y%m%d%H%M%S")+".pdf"
#
# reader = document
# number_of_pages = len(reader.pages)
# k=""
# for i in range(0,number_of_pages):
#     page = reader.pages[i]
#     text = page.extract_text()
#     k=k+text
#
# print(k)

import docx

doc=docx.Document("C:\\Users\\Reema\\PycharmProjects\\nlp\\media\\20230311134840Word.docx")

for i in doc.paragraphs:
    print(i.text)