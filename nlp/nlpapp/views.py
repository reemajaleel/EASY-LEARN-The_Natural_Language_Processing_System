import io

from django.core.files.storage import FileSystemStorage
from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
from nltk import sent_tokenize
from wikipedia import wikipedia
from yake import yake
from fpdf import FPDF

# import os
# os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'


from nlpapp.models import *

import pandas

from pypdf import PdfReader


# Create your views here.
# from nlpapp.pipelines import pipeline
from nlpapp.pipelines import pipeline


def home(request):
    return render(request,"home.html")
def loginfun(request):
    return render(request,"index.html")
def login_post(request):
    username=request.POST["textfield"]
    password=request.POST["textfield2"]
    res=login.objects.filter(username=username, password=password)
    if res.exists():
        res1 = login.objects.get(username=username, password=password)
        request.session['lid']=res1.id
        if res1.type =='admin' :
            return redirect('/myapp/dashboard1/')
        elif res1.type == 'user':
            return redirect('/myapp/dashboard/')
        else:
            return redirect('/myapp/login/')
    else:
        return redirect('/myapp/login/')


def adminviewusers(request):
    return render(request,"view registered users.html")




def userregister(request):
    return render(request,"signupindex.html")

def userregister_post(request):
    firstname=request.POST["textfield"]
    lastname=request.POST["textfield2"]
    dateofbirth=request.POST["textfield3"]
    email=request.POST["textfield4"]
    phonenumber=request.POST["textfield5"]
    password=request.POST["textfield6"]

    if user.objects.filter(email=email, phone=phonenumber).exists():
        return HttpResponse("<script>alert('Already a user');window.location='/myapp/login';</script>")


    lobj=login()
    lobj.username=email
    lobj.password=password
    lobj.type="user"
    lobj.save()

    robj=user()
    robj.fname=firstname
    robj.lname=lastname
    robj.dob=dateofbirth
    robj.email=email
    robj.phone=phonenumber
    robj.LOGIN=lobj
    robj.save()
    return render(request,"index.html")

def userviewprofile(request):
    lid = request.session['lid']
    res = user.objects.get(LOGIN_id=lid)
    return render(request,"userviewprofile1.html",{'data':res})

def editprofile(request):
    lid = request.session['lid']
    res = user.objects.get(LOGIN_id=lid)
    return render(request,"editprofile1.html",{'data':res})

def editprofile_post(request):
    firstname = request.POST["textfield"]
    lastname = request.POST["textfield2"]
    dateofbirth = request.POST["textfield3"]
    email = request.POST["textfield4"]
    phonenumber = request.POST["textfield5"]
    lid = request.session['lid']
    res = user.objects.filter(LOGIN_id=lid).update(fname=firstname,lname=lastname,dob=dateofbirth,email=email,phone=phonenumber)
    res = user.objects.get(LOGIN_id=lid)
    return render(request,"userviewprofile1.html",{'data':res})

def useruploadtext(request):
    return render(request,"uploadtext1.html")
def useruploadtext_post(request):
    text=request.POST["textarea"]
    return render(request,"upload text.html")
def useruploaddocument(request):
    return render(request,"uploaddocument.html")
def useruploaddocument_post(request):
    choose=request.POST["select"]
    return render(request,"upload pdf or word.html")
def useruploadvoice(request):
    return render(request,"upload voice.html")
def useruploadvoice_post(request):
    text=request.POST["textarea"]
    return render(request,"upload voice.html")
def userwiki(request):
    return render(request,"keyword and wiki content extraction.html")
def usersummary(request):
    return render(request,"summary.html")
def userqa(request):
    return render(request,"questions and answers.html")
def index(request):
    return render(request,"index.html")
def signupindex(request):
    return render(request,"signupindex.html")
def regusers(request):
    res = user.objects.all()
    return render(request,"regusers.html",{'data':res})
def uploaddocument(request):
    return render(request,"uploaddocument.html")




def uploaddocument_post(request):
    # import pyPdf
    from pypdf import PdfReader

    photo=request.FILES['document']
    from datetime import datetime
    date=datetime.now().strftime("%Y%m%d%H%M%S") +photo.name
    fs=FileSystemStorage()
    fn=fs.save(date,photo)
    path=fs.url(date)
    ext=photo.name.split(".")

    s=""
    print(ext,"fgdhfdfgdhfd")

    if ext[1] =='pdf':
        reader = PdfReader("C:\\Users\\Reema\\PycharmProjects\\nlp\\media\\"+date)
        number_of_pages = len(reader.pages)
        k = ""
        for i in range(0, number_of_pages):
            page = reader.pages[i]
            text = page.extract_text()
            k = k + text
        print(k)
        s=k


    if ext[1] == 'docx':
        import docx

        doc = docx.Document("C:\\Users\\Reema\\PycharmProjects\\nlp\\media\\"+date)

        for i in doc.paragraphs:
            s=s+i.text
            print(i.text)
    # request.session['extrs']=s

    return render(request,'extractedtext2.html',{'s':s})


def extractedtext1(request):
    s=request.POST['textarea']
    return render(request, 'extractedtext1.html', {'s': s})

def extractedtext2(request):
    s=request.POST['textarea']
    return render(request, 'extractedtext.html', {'s': s})


def dashboard(request):
    return render(request,"dashboard.html")
def dashboard1(request):
    return render(request,"dashboard1.html")

def uploadtext1(request):
    return render(request,"uploadtext1.html")

def uploadtext1_post(request):
    return HttpResponse("ok")
    return render(request,"extractedtext.html")

def uploadvoice1(request):
    return render(request,"uploadvoice1.html")

def uploadvoice1_post(request):
    return HttpResponse("ok")
    return render(request, "uploadvoice1.html")

def extractedtext(request):
    return render(request,"extractedtext.html")

def noteinput(request):
    return render(request, 'note_index.html')

# def notpost(request):
#     name = request.POST['textarea']
#     return render(request, 'extractedtext.html','s':name)
#



def keywordandwiki1(request):
    message = request.POST['wc']
    language = "en"
    max_ngram_size = 1
    deduplication_threshold = 0.9
    numOfKeywords = 5
    custom_kw_extractor = yake.KeywordExtractor(lan=language, n=max_ngram_size,
                                                     dedupLim=deduplication_threshold, top=numOfKeywords,
                                                     features=None)
    keywords = custom_kw_extractor.extract_keywords(message)
    # wikipedia.set_lang("hi")
    wikicontent=[]



    g=message

    for i in keywords:
        print(i[0])
        try:
            s = wikipedia.summary(str(i[0]))
            # print(s)

            s=str(s).replace("'","")
            s=str(s).replace('"',"")
            s=str(s).replace(';',"")

            # g= g +s
            wikicontent.append(s)
        except:
            wikicontent.append("")
            print("error")


    # print(wikicontent)
    # print(keywords)
    kw=[]
    wc=[]
    for i in keywords:
        kw.append(i[0])
    for i in wikicontent:
        wc.append(i)

    request.session["keywords"] = kw
    request.session["wikicontent"] = wc

    return render(request,"keywordandwiki1.html",{'kw':kw,'wc':wc})



def summary1(request):
    import transformers
    from transformers import pipeline
    from transformers import PreTrainedTokenizer
    from transformers import BertTokenizer
    from trainer import Trainer
    # from torch.optim.lr_scheduler import SAVE_STATE_WARNING

    to_tokenize = request.POST['ts']
    summarizer = pipeline("summarization")
    summarized = summarizer(to_tokenize, min_length=75, max_length=300)

    print(summarized)


    summary_text = [summary['summary_text'] for summary in summarized]

    print(summary_text)

    request.session["summary_text"] = summary_text

    return render(request,"summary1.html",{'summarized':  summary_text })





def qa1(request):

    import os
    import docx
    g=request.POST['qa']
    qa=pipeline('question-generation')
    s=qa(g)
    d=""
    m=1
    print(s)
    question=[]
    answer=[]
    for i in s:
        question.append(i["question"])
        answer.append(i["answer"])
        d = d + str(m) + "." + i['question'] + "\n" + "Answer: " + i['answer'] + "\n\n\n\n"
        m=m+1
        print(d)

    request.session["d"] = d
    request.session["question"]=question
    request.session["answer"]=answer
    return render(request, "qa1.html", {'d': s})



#
# def generatevoice(request):
#     text_val = request.POST['voice']
#     # question = request.session["question"]
#     # answer = request.session["answer"]
#     # print(type(answer), type(question))
#     d = request.session["d"]
#     print(d)
#
#
#     # Import the gTTS module for text to speech conversion
#     from gtts import gTTS
#
#     # This module is imported so that we can play the converted audio
#     from playsound import playsound
#
#     # Create a list of QA pairs with the question before the answer
#     # qa_pairs = [f"{i+1}. {q} {a}" for i, (q, a) in enumerate(zip(question, answer))]
#
#     qa_pairs= ''.join([f"{dqa}" for dqa in zip(d)])
#
#     # Concatenate the text input and the QA pairs
#     # full_text = text_val + " " + " ".join(qa_pairs)
#
#     # Create a gTTS object for the full text
#     language = 'en'
#     gtts_obj = gTTS(text=text_val, lang=language, slow=False)
#
#     # Here we are saving the transformed audio in a mp3 file named qavoice.mp3
#     gtts_obj.save("C://Users//Reema//PycharmProjects//nlp//media//qavoice.mp3")
#     gtts_obj.save("C://Users//Reema//Desktop//qavoice.mp3")
#
#     # Play the qavoice.mp3 file
#     # playsound("C://Users//Reema//PycharmProjects//nlp//media//qavoice.mp3")
#
#
#     return render(request, "success1.html")

def generatevoice(request):
    text_val = request.POST['voice']
    d = request.session["d"]
    print(d)

    # Import the gTTS module for text to speech conversion
    from gtts import gTTS

    # This module is imported so that we can play the converted audio
    from playsound import playsound

    # Create a list of QA pairs with the question before the answer
    qa_pairs= ''.join([f"{dqa}" for dqa in zip(d)])

    # Create a gTTS object for the full text
    language = 'en'
    gtts_obj = gTTS(text=d, lang=language, slow=False)

    # Here we are saving the transformed audio in a mp3 file named qavoice.mp3
    # gtts_obj.save("C://Users//Reema//PycharmProjects//nlp//media//qavoice.mp3")
    gtts_obj.save("C://Users//Reema//Desktop//qavoice.mp3")

    # Play the qavoice.mp3 file
    # playsound("C://Users//Reema//PycharmProjects//nlp//media//qavoice.mp3")


    return render(request, "success1.html")

#
# def stopaudio(request):
#     st=request.POST['stop']
#     # Retrieve the player object from the session variable
#     p = request.session['p']
#
#     if p:
#         # Stop the audio playback and remove the session variable
#         p.stop()
#         del request.session["p"]
#
#     return HttpResponse("Audio playback stopped")


def generatevoice1(request):
    text_val = request.POST['voice1']
    kw = request.session["keywords"]
    wc = request.session["wikicontent"]
    print(kw)
    print(wc)

    s = ' '.join([f"{k} {w}" for k, w in zip(kw, wc)])

    # Import the gTTS module for text
    # to speech conversion
    from gtts import gTTS

    # This module is imported so that we can
    # play the converted audio

    from playsound import playsound

    # Here are converting in English Language
    language = 'en'


    # Passing the text and language to the engine,
    # here we have assign slow=False. Which denotes
    # the module that the transformed audio should
    # have a high speed
    obj = gTTS(text=s, lang=language, slow=False)

    # Here we are saving the transformed audio in a mp3 file named
    # exam.mp3
    # obj.save("C://Users//Reema//PycharmProjects//nlp//media//kw.mp3")
    obj.save("C://Users//Reema//Desktop//kw.mp3")

    # Play the exam.mp3 file
    # playsound("C://Users//Reema//PycharmProjects//nlp//media//kw.mp3")

    return render(request, "success1.html")


def generatepdf(request):
    from datetime import datetime
    g=request.POST['pdf']
    question=request.session["question"]
    answer=request.session["answer"]
    print(type(answer),type(question))
    # qa = pipeline('question-generation')

    s = g
    d = ""
    m = 1
    print(s)

    from fpdf import FPDF

    pdf = FPDF()

    pdf.add_page()

    pdf.set_font("Arial", size=10)
    f = g
    j=0
    print(type(s))
    for i in range(len(question)):
        print(i)
        print(type(i))
        j = j + 1

        pdf.cell(0, 10, f"{str(j)}. {question[i]}", ln=1)
        pdf.cell(0, 10, f"Answer: {answer[i]}", ln=1)
        pdf.cell(0, 10, '', ln=1)

    # pdf.output("C://Users//Reema//PycharmProjects//nlp//media//pp2.pdf")
    pdf.output("C://Users//Reema//Desktop//pp2.pdf")
    return render(request, "success.html")



def generatepdf1(request):
    g = request.POST['pdf1']
    kw = request.session["keywords"]
    wc = request.session["wikicontent"]
    print(kw)
    print(wc)

    s = g

    from fpdf import FPDF

    pdf = FPDF()

    pdf.add_page()

    pdf.set_font("Arial", size=15)

    pdf.multi_cell(0, 10, f"Keywords: {kw}" )
    pdf.multi_cell(0, 10, f"Wikicontent: {wc}")

    # pdf.output("C://Users//Reema//PycharmProjects//nlp//media//pp3.pdf")
    pdf.output("C://Users//Reema//Desktop//pp3.pdf")

    return render(request, "success.html")

def generatepdf2(request):
    g=request.POST['pdf2']
    summary_text = request.session["summary_text"]
    print(summary_text)

    s = g


    from fpdf import FPDF

    pdf = FPDF()

    pdf.add_page()

    pdf.set_font("Arial", size=15)

    pdf.multi_cell(0, 10, f"Summary: {summary_text}")

    # pdf.output("C://Users//Reema//PycharmProjects//nlp//media//pp4.pdf")
    pdf.output("C://Users//Reema//Desktop//pp4.pdf")

    return render(request, "success.html")

#               voice                   #


def generatevoice2(request):
    text_val = request.POST['voice2']
    summary_text = request.session["summary_text"]
    print(summary_text)

    # join the strings in the list comprehension into a single string
    s = ''.join([f"{su}" for su in zip(summary_text)])

    from gtts import gTTS
    from playsound import playsound
    language = 'en'

    obj = gTTS(text=s, lang=language, slow=False)

    # obj.save("C://Users//Reema//PycharmProjects//nlp//media//summary.mp3")
    obj.save("C://Users//Reema//Desktop//summary.mp3")

    # Play the exam.mp3 file
    # playsound("C://Users//Reema//PycharmProjects//nlp//media//summary.mp3")

    # return a response object
    return render(request, "success1.html")






def generatedocument(request):
    from docx import Document

    g = request.POST['document']
    question = request.session["question"]
    answer = request.session["answer"]
    print(type(answer), type(question))

    document = Document()

    # Add a heading
    document.add_heading('Questions and Answers', 0)

    # Loop through the questions and answers and add them to the document
    for i in range(len(question)):
        p = document.add_paragraph(f"{str(i+1)}. {question[i]}")
        p.add_run().text = f" Answer: {answer[i]}"
        document.add_paragraph()

    # Save the document to a file
    # document.save('C://Users//Reema//PycharmProjects//nlp//media//pp2.docx')
    document.save('C://Users//Reema//Desktop//pp2.docx')

    return render(request, "success.html")



def generatedocument1(request):
    from docx import Document

    g = request.POST['document1']
    kw = request.session["keywords"]
    wc = request.session["wikicontent"]
    print(kw)
    print(wc)

    document = Document()

    # Add a heading
    document.add_heading('Keywords and Wikicontent', 0)

    # Add the keywords and wikicontent to the document
    p = document.add_paragraph(f"Keywords: {kw}")
    document.add_paragraph()
    p = document.add_paragraph(f"Wikicontent: {wc}")

    # Save the document to a file
    # document.save('C://Users//Reema//PycharmProjects//nlp//media//pp3.docx')
    document.save("C://Users//Reema//Desktop//pp3.docx")

    return render(request, "success.html")




def generatedocument2(request):
    from docx import Document

    text_val = request.POST['document2']
    summary_text = request.session["summary_text"]
    print(summary_text)

    # join the strings in the list comprehension into a single string
    s = ''.join([f"{su}" for su in zip(summary_text)])

    document = Document()

    # Add a heading
    document.add_heading('Summary', 0)

    # Add the summary text to the document
    p = document.add_paragraph(s)

    # Save the document to a file
    # document.save('C://Users//Reema//PycharmProjects//nlp//media//summary.docx')
    document.save("C://Users//Reema//Desktop//summary.docx")

    return render(request, "success1.html")


#
# def adminviewcomplaint(request):
#     return render(request,"view complaint.html")


# def adminsendreply(request):
#     return render(request,"send reply.html")
# def adminsendreply_post(request):
#     name=request.POST["textfield"]
#     complaint=request.POST["textarea"]
#     reply=request.POST["textarea2"]
#     return render(request,"send reply.html")

def viewcomplaint(request):
    # res4 = user.objects.all()
    res5=Complaint.objects.all()
    # data1 = {'res4': res4, 'res5': res5}
    return render(request, "view complaint.html", {'datas' : res5})


def search_post(request):
    fromdate=request.POST['textfield']
    todate=request.POST['textfield2']
    res6 = Complaint.objects.filter(date__range=(fromdate,todate))
    return render(request, "view complaint.html", {'datas' : res6})



def success2(request):
    return render(request, "success2.html")

def sendcomplaint(request):
    return render(request, 'sendcomplaint.html')


def sendcomplaint_post(request):
    from datetime import datetime
    lid = request.session['lid']
    c=request.POST['complaint']
    cobj=Complaint()
    cobj.complaint=c
    cobj.date=datetime.now().date()
    cobj.reply='pending'
    cobj.status='pending'
    cobj.USERID=user.objects.get(LOGIN_id=lid)
    cobj.save()
    return render(request, 'success2.html')


def sendreply(request,cid):
    res = Complaint.objects.get(id=cid)
    return render(request, "sendreply1.html",{"data":res})


def sendreply_post(request):
    r=request.POST['textarea2']
    id=request.POST['id']
    rply = Complaint.objects.filter(id=id).update(reply=r, status="Replied")
    return redirect('/myapp/viewcomplaint/')


def viewreply(request):
    res6 = Complaint.objects.filter(USERID=user.objects.get(LOGIN_id=request.session['lid']))
    return render(request, "viewreply.html", {"datareply": res6})



def generatepdff(request):
    pd=request.POST['ppddff']
    # summary_text = request.session["summary_text"]
    # print(summary_text)


    from fpdf import FPDF

    pdf = FPDF()

    pdf.add_page()

    pdf.set_font("Arial", size=15)

    pdf.multi_cell(0, 10, f"{pd}")

    # pdf.output("C://Users//Reema//PycharmProjects//nlp//media//pp4.pdf")
    pdf.output("C://Users//Reema//Desktop//text.pdf")

    return render(request, "success.html")




def generatedocumentt(request):
    from docx import Document

    text_val = request.POST['ddoocc']
    # summary_text = request.session["summary_text"]
    # print(summary_text)

    # join the strings in the list comprehension into a single string
    # s = ''.join([f"{te}" for te in zip(text_val)])

    document = Document()

    # Add a heading
    # document.add_heading('Text', 0)

    # Add the summary text to the document
    p = document.add_paragraph(text_val)

    # Save the document to a file
    # document.save('C://Users//Reema//PycharmProjects//nlp//media//summary.docx')
    document.save("C://Users//Reema//Desktop//text.docx")

    return render(request, "success1.html")




def generatevoicee(request):
    text_val = request.POST['vvooiiccee']
    # d = request.session["d"]
    # print(d)

    # Import the gTTS module for text to speech conversion
    from gtts import gTTS

    # This module is imported so that we can play the converted audio
    from playsound import playsound

    # Create a list of QA pairs with the question before the answer
    qa_pairs= ''.join([f"{te}" for te in zip(text_val)])

    # Create a gTTS object for the full text
    language = 'en'
    gtts_obj = gTTS(text=text_val, lang=language, slow=False)

    # Here we are saving the transformed audio in a mp3 file named qavoice.mp3
    # gtts_obj.save("C://Users//Reema//PycharmProjects//nlp//media//qavoice.mp3")
    gtts_obj.save("C://Users//Reema//Desktop//text.mp3")

    # Play the qavoice.mp3 file
    # playsound("C://Users//Reema//PycharmProjects//nlp//media//qavoice.mp3")


    return render(request, "success1.html")
